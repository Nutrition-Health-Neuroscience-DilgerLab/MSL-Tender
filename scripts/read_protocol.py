"""
Extract text from the naming protocol Word document
"""
from docx import Document

doc_path = r"c:\Users\rndpi\Documents\Coding Projects\MSL-tender\UI 23-16 Chop naming and data pairing protocol v1.docx"

doc = Document(doc_path)

print("="*80)
print("CHOP NAMING AND DATA PAIRING PROTOCOL")
print("="*80)
print()

for paragraph in doc.paragraphs:
    if paragraph.text.strip():
        print(paragraph.text)
        print()

# Also check for tables
if doc.tables:
    print("\n" + "="*80)
    print("TABLES IN DOCUMENT")
    print("="*80)
    for i, table in enumerate(doc.tables):
        print(f"\nTable {i+1}:")
        for row in table.rows:
            print(" | ".join(cell.text.strip() for cell in row.cells))
